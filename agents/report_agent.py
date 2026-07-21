"""Agent 3: reads all per-API markdown summaries and writes the final consolidated .docx report."""

import json
import re
import sys
from datetime import date
from pathlib import Path
from typing import Any

if __name__ == "__main__" and __package__ is None:
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from pydantic import BaseModel, ValidationError

from agents.common import PROJECT_ROOT, get_llm, setup_logging

logger = setup_logging(__name__)

SUMMARIES_DIR: Path = PROJECT_ROOT / "results" / "summaries"
FINAL_REPORT_DIR: Path = PROJECT_ROOT / "final_report"

COLOR_HEALTHY = RGBColor(0x00, 0xB0, 0x50)
COLOR_DEGRADING = RGBColor(0xFF, 0x99, 0x00)
COLOR_CRITICAL = RGBColor(0xFF, 0x00, 0x00)

BODY_FONT = "Calibri"
BODY_SIZE = Pt(11)
HEADING_SIZE = Pt(14)

EXEC_SUMMARY_PROMPT = """You are a performance engineering lead writing an executive summary \
for a stress test report covering {count} APIs. Respond with ONLY a JSON object matching \
exactly this schema (no markdown fences, no extra text):

{{"summary": "<2-3 sentence overview comparing overall resilience across the APIs>", \
"most_resilient_api": "<name of the API that held up best under load>", \
"most_fragile_api": "<name of the API that degraded earliest>", \
"overall_recommendation": "<one sentence overall system recommendation>"}}

Per-API summary (name, safe limit, breaking point, max error %):
{overview}
"""


class ExecutiveSummary(BaseModel):
    """Schema-validated LLM output for the report's executive summary."""

    summary: str
    most_resilient_api: str
    most_fragile_api: str
    overall_recommendation: str


def _strip_code_fences(text: str) -> str:
    """Strip a leading/trailing ```json ... ``` fence from an LLM response, if present.

    Args:
        text: Raw LLM response text.

    Returns:
        Text with any surrounding markdown code fence removed.
    """
    stripped = text.strip()
    match = re.match(r"^```(?:json)?\s*(.*?)\s*```$", stripped, re.DOTALL)
    return match.group(1) if match else stripped


def _status_color(status: str) -> RGBColor:
    """Map a status label (as emitted by the summariser agent) to a docx font color.

    Args:
        status: Status string, e.g. "✅ Healthy", "⚠️ Degrading", "💀 Critical".

    Returns:
        RGBColor for the matching health tier; defaults to critical red if unrecognised.
    """
    if "Healthy" in status:
        return COLOR_HEALTHY
    if "Degrading" in status:
        return COLOR_DEGRADING
    return COLOR_CRITICAL


def parse_summary_md(md_text: str) -> dict[str, Any]:
    """Parse a markdown summary file produced by summariser_agent into structured data.

    Args:
        md_text: Full contents of a results/summaries/{api_name}.md file.

    Returns:
        Dict with name, url, method, test_date, verdict, rows, regression_flags,
        bottleneck_hypothesis, safe_limit, breaking_point, breaking_pct, and recommendation.
    """
    name_match = re.search(r"^# API Stress Test Summary: (.+)$", md_text, re.MULTILINE)
    url_match = re.search(r"\*\*URL:\*\* (.+)", md_text)
    method_match = re.search(r"\*\*Method:\*\* (.+)", md_text)
    date_match = re.search(r"\*\*Test Date:\*\* (.+)", md_text)
    verdict_match = re.search(r"\*\*Verdict:\*\* (.+)", md_text)
    safe_match = re.search(r"\*\*Safe limit:\*\* (\S+)", md_text)
    breaking_match = re.search(
        r"\*\*Breaking point:\*\* (\S+) concurrent users \(([\d.]+%|N/A)", md_text
    )
    regression_match = re.search(
        r"## Regression vs Previous Run\n\n(.+?)\n\n## Bottleneck Hypothesis", md_text, re.DOTALL
    )
    bottleneck_match = re.search(
        r"## Bottleneck Hypothesis\n\n(.+?)\n\n## Breaking Point", md_text, re.DOTALL
    )
    recommendation_match = re.search(
        r"## Recommendation\n\n(.+?)\s*$", md_text, re.DOTALL
    )

    rows: list[dict[str, str]] = []
    for line in md_text.splitlines():
        if (
            line.startswith("|")
            and not line.startswith("| -")
            and "Threads" not in line
        ):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if len(cells) == 10 and cells[0].isdigit():
                rows.append(
                    {
                        "threads": cells[0],
                        "requests": cells[1],
                        "avg": cells[2],
                        "min": cells[3],
                        "max": cells[4],
                        "p95": cells[5],
                        "p99": cells[6],
                        "throughput": cells[7],
                        "error_pct": cells[8],
                        "status": cells[9],
                    }
                )

    regression_text = regression_match.group(1).strip() if regression_match else ""
    regression_flags = (
        [line.lstrip("- ").strip() for line in regression_text.splitlines() if line.strip().startswith("-")]
        if regression_text
        else []
    )

    return {
        "name": name_match.group(1).strip() if name_match else "unknown",
        "url": url_match.group(1).strip() if url_match else "",
        "method": method_match.group(1).strip() if method_match else "",
        "test_date": date_match.group(1).strip() if date_match else "",
        "verdict": verdict_match.group(1).strip() if verdict_match else "unknown",
        "rows": rows,
        "regression_flags": regression_flags,
        "bottleneck_hypothesis": bottleneck_match.group(1).strip() if bottleneck_match else "",
        "safe_limit": safe_match.group(1).strip() if safe_match else "N/A",
        "breaking_point": (
            breaking_match.group(1).strip() if breaking_match else "Not reached"
        ),
        "breaking_pct": breaking_match.group(2).strip() if breaking_match else "N/A",
        "recommendation": (
            recommendation_match.group(1).strip() if recommendation_match else ""
        ),
    }


def load_all_summaries(api_names: set[str] | None = None) -> list[dict[str, Any]]:
    """Load and parse markdown summary files in results/summaries/.

    Args:
        api_names: If given, only summaries whose parsed name is in this set are
            included — filters out stale files left over from APIs that were
            renamed or removed from config.yaml in a later run. If None, every
            .md file present is included (legacy behavior).

    Returns:
        List of parsed summary dicts, one per API, in filename-sorted order.
    """
    summaries = []
    if not SUMMARIES_DIR.exists():
        return summaries
    for md_path in sorted(SUMMARIES_DIR.glob("*.md")):
        try:
            with open(md_path, encoding="utf-8") as f:
                parsed = parse_summary_md(f.read())
        except OSError as exc:
            logger.warning("Could not read summary %s: %s", md_path, exc)
            continue

        if api_names is not None and parsed["name"] not in api_names:
            logger.info(
                "Skipping stale summary %s (API not in current config.yaml)", md_path.name
            )
            continue

        summaries.append(parsed)
    return summaries


def _max_error_pct(summary: dict[str, Any]) -> float:
    """Compute the maximum error % observed across a summary's rows.

    Args:
        summary: Parsed summary dict as returned by parse_summary_md.

    Returns:
        Maximum error percentage as a float, or 0.0 if no rows are present.
    """
    values = []
    for row in summary["rows"]:
        try:
            values.append(float(row["error_pct"].rstrip("%")))
        except ValueError:
            continue
    return max(values) if values else 0.0


def report_title_slug(config: dict[str, Any]) -> str:
    """Derive a filesystem-safe slug from settings.report_title for use in output filenames.

    Args:
        config: Parsed config.yaml dictionary.

    Returns:
        The report title with non-alphanumeric runs collapsed to underscores.
    """
    title = config["settings"].get("report_title", "API Stress Test Report")
    return re.sub(r"[^A-Za-z0-9]+", "_", title).strip("_")


def build_report_filename(config: dict[str, Any]) -> str:
    """Build the dated output filename for the final report from settings.report_title.

    Args:
        config: Parsed config.yaml dictionary.

    Returns:
        Filename in the form "{Report_Title_Slug}_{date}.docx".
    """
    return f"{report_title_slug(config)}_{date.today().isoformat()}.docx"


def _overall_status(max_error_pct: float) -> str:
    """Map a max error % to an overall status label using CLAUDE.md thresholds.

    Args:
        max_error_pct: Maximum error percentage observed for an API.

    Returns:
        One of "Healthy", "Degrading", "Critical".
    """
    if max_error_pct < 1.0:
        return "Healthy"
    if max_error_pct <= 10.0:
        return "Degrading"
    return "Critical"


def _fallback_executive_summary(summaries: list[dict[str, Any]]) -> dict[str, str]:
    """Produce a template-based (non-LLM) executive summary.

    Args:
        summaries: List of parsed per-API summary dicts.

    Returns:
        Dict with "summary", "most_resilient_api", "most_fragile_api", "overall_recommendation".
    """
    ranked = sorted(summaries, key=_max_error_pct)
    most_resilient = ranked[0]["name"] if ranked else "N/A"
    most_fragile = ranked[-1]["name"] if ranked else "N/A"
    return {
        "summary": (
            f"This report covers {len(summaries)} APIs tested under staircase load. "
            f"{most_resilient} showed the strongest resilience under load, while {most_fragile} "
            f"degraded earliest."
        ),
        "most_resilient_api": most_resilient,
        "most_fragile_api": most_fragile,
        "overall_recommendation": (
            "Cap production concurrency at each API's individually determined safe limit, "
            "with the most fragile services prioritised for remediation."
        ),
    }


def generate_executive_summary(config: dict[str, Any], summaries: list[dict[str, Any]]) -> dict[str, str]:
    """Ask the configured LLM for a schema-validated executive summary, falling back to a template.

    Args:
        config: Parsed config.yaml dictionary.
        summaries: List of parsed per-API summary dicts.

    Returns:
        Dict with "summary", "most_resilient_api", "most_fragile_api", "overall_recommendation".
    """
    overview_lines = []
    for s in summaries:
        overview_lines.append(
            f"- {s['name']}: safe limit={s['safe_limit']}, breaking point={s['breaking_point']}, "
            f"max error={_max_error_pct(s):.2f}%"
        )
    overview = "\n".join(overview_lines)

    try:
        llm = get_llm(config)
        prompt = EXEC_SUMMARY_PROMPT.format(count=len(summaries), overview=overview)
        response = llm.invoke(prompt)
        text = response.content if hasattr(response, "content") else str(response)
        parsed = ExecutiveSummary.model_validate_json(_strip_code_fences(text))
        return parsed.model_dump()
    except (ValidationError, json.JSONDecodeError) as exc:
        logger.warning("Executive summary failed schema validation: %s — using template fallback", exc)
    except Exception as exc:  # noqa: BLE001 - any LLM failure must fall back, never crash
        logger.warning("Executive summary LLM call failed: %s — using template fallback", exc)

    return _fallback_executive_summary(summaries)


def _set_body_style(document: Document) -> None:
    """Set the default document font to Calibri 11pt per CLAUDE.md styling rules.

    Args:
        document: The python-docx Document to style.
    """
    style = document.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE


def _add_heading(document: Document, text: str, level: int) -> None:
    """Add a heading paragraph styled with the Calibri font and 14pt size.

    Args:
        document: The python-docx Document to append to.
        text: Heading text.
        level: Heading level (1 for section titles, 2 for subsections).
    """
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = BODY_FONT
        run.font.size = HEADING_SIZE


def _add_results_table(document: Document, rows: list[dict[str, str]]) -> None:
    """Add a color-coded results table for one API's staircase results.

    Args:
        document: The python-docx Document to append to.
        rows: Parsed table rows (threads, requests, avg, min, max, p95, p99, throughput,
            error_pct, status).
    """
    table = document.add_table(rows=1, cols=10)
    table.style = "Light Grid Accent 1"
    headers = [
        "Threads",
        "Requests",
        "Avg (ms)",
        "Min (ms)",
        "Max (ms)",
        "P95 (ms)",
        "P99 (ms)",
        "Throughput (rps)",
        "Error %",
        "Status",
    ]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header

    for row in rows:
        cells = table.add_row().cells
        values = [
            row["threads"],
            row["requests"],
            row["avg"],
            row["min"],
            row["max"],
            row["p95"],
            row["p99"],
            row["throughput"],
            row["error_pct"],
            row["status"],
        ]
        color = _status_color(row["status"])
        for cell, value in zip(cells, values):
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = color


def _add_footer(document: Document) -> None:
    """Add the standard footer text to every section of the document.

    Args:
        document: The python-docx Document to add a footer to.
    """
    for section in document.sections:
        footer = section.footer
        paragraph = (
            footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        )
        paragraph.text = "Generated by JMeter Agent"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = BODY_FONT
            run.font.size = Pt(9)


def build_report(config: dict[str, Any], summaries: list[dict[str, Any]]) -> Document:
    """Assemble the full consolidated .docx report document.

    Args:
        config: Parsed config.yaml dictionary.
        summaries: List of parsed per-API summary dicts.

    Returns:
        The assembled python-docx Document, ready to be saved.
    """
    document = Document()
    _set_body_style(document)

    # Cover page
    report_title = config["settings"].get("report_title", "API Stress Test Report")
    title = document.add_heading(report_title, level=0)
    for run in title.runs:
        run.font.name = BODY_FONT
    subtitle = document.add_paragraph(f"Test Date: {date.today().isoformat()}")
    subtitle.add_run("\nEnvironment: Mac (Apple Silicon), JMeter 5.6.3")
    document.add_page_break()

    # Table of contents
    _add_heading(document, "Table of Contents", level=1)
    document.add_paragraph("1. Executive Summary")
    for i, s in enumerate(summaries, start=2):
        document.add_paragraph(f"{i}. {s['name']}")
    document.add_paragraph(f"{len(summaries) + 2}. Overall System Health Summary")
    document.add_page_break()

    # Executive summary
    exec_summary = generate_executive_summary(config, summaries)
    _add_heading(document, "Executive Summary", level=1)
    document.add_paragraph(exec_summary["summary"])
    document.add_paragraph(f"Most resilient: {exec_summary['most_resilient_api']}")
    document.add_paragraph(f"Most fragile: {exec_summary['most_fragile_api']}")
    document.add_paragraph(f"Overall recommendation: {exec_summary['overall_recommendation']}")
    document.add_page_break()

    # Per-API sections
    for s in summaries:
        _add_heading(document, s["name"], level=1)
        document.add_paragraph(f"URL: {s['url']}")
        document.add_paragraph(f"Method: {s['method']}")
        document.add_paragraph(f"Test Date: {s['test_date']}")
        document.add_paragraph(f"Verdict: {s['verdict']}")

        _add_heading(document, "Results", level=2)
        _add_results_table(document, s["rows"])

        _add_heading(document, "Regression vs Previous Run", level=2)
        if s["regression_flags"]:
            for flag in s["regression_flags"]:
                document.add_paragraph(flag, style="List Bullet")
        else:
            document.add_paragraph("No regressions detected vs the previous run.")

        _add_heading(document, "Bottleneck Hypothesis", level=2)
        document.add_paragraph(s["bottleneck_hypothesis"])

        _add_heading(document, "Breaking Point", level=2)
        document.add_paragraph(f"Safe limit: {s['safe_limit']} concurrent users")
        document.add_paragraph(
            f"Breaking point: {s['breaking_point']} concurrent users ({s['breaking_pct']} error rate)"
        )

        _add_heading(document, "Recommendation", level=2)
        document.add_paragraph(s["recommendation"])
        document.add_page_break()

    # Overall system health summary
    _add_heading(document, "Overall System Health Summary", level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for cell, header in zip(
        table.rows[0].cells,
        ["API Name", "Safe Limit", "Breaking Point", "Max Error %", "Status"],
    ):
        cell.text = header
    for s in summaries:
        max_pct = _max_error_pct(s)
        status = _overall_status(max_pct)
        cells = table.add_row().cells
        for cell, value in zip(
            cells,
            [
                s["name"],
                s["safe_limit"],
                s["breaking_point"],
                f"{max_pct:.2f}%",
                status,
            ],
        ):
            cell.text = value

    _add_footer(document)
    return document


def main() -> None:
    """Entry point: read all summaries and write the final consolidated .docx report."""
    from agents.common import load_config

    config = load_config()
    configured_api_names = {api["name"] for api in config["apis"]}
    summaries = load_all_summaries(configured_api_names)

    if not summaries:
        logger.warning(
            "No summaries found in %s, skipping report generation", SUMMARIES_DIR
        )
        return

    document = build_report(config, summaries)

    FINAL_REPORT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = FINAL_REPORT_DIR / build_report_filename(config)
    document.save(str(output_path))
    logger.info("Final report saved to %s", output_path)


if __name__ == "__main__":
    main()
